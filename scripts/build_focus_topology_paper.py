from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "专注音乐多尺度有向拓扑指纹_完整论文.md"
OUTPUT = ROOT / "output" / "专注音乐多尺度有向拓扑指纹_完整论文.docx"

INK = RGBColor(0x20, 0x38, 0x64)
BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
MUTED = RGBColor(0x66, 0x66, 0x66)
BLACK = RGBColor(0, 0, 0)
TABLE_FILL = "F4F6F9"
TABLE_BORDER = "B8C4D1"
CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(
    run,
    *,
    latin: str = "Calibri",
    east_asia: str = "SimSun",
    size: float | None = None,
    color: RGBColor | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:cs"), latin)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, *, latin: str, east_asia: str, size: float, color=BLACK) -> None:
    style.font.name = latin
    style.font.size = Pt(size)
    style.font.color.rgb = color
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:cs"), latin)


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    set_run_font(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    set_style_font(normal, latin="Calibri", east_asia="SimSun", size=11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.widow_control = True

    title = doc.styles["Title"]
    set_style_font(title, latin="Calibri", east_asia="Microsoft YaHei", size=28, color=INK)
    title.font.bold = True
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(12)

    subtitle = doc.styles["Subtitle"]
    set_style_font(subtitle, latin="Calibri", east_asia="Microsoft YaHei", size=14, color=DARK_BLUE)
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(18)

    h1 = doc.styles["Heading 1"]
    set_style_font(h1, latin="Calibri", east_asia="Microsoft YaHei", size=16, color=BLUE)
    h1.font.bold = True
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.line_spacing = 1.0
    h1.paragraph_format.keep_with_next = True
    h1.paragraph_format.widow_control = True

    h2 = doc.styles["Heading 2"]
    set_style_font(h2, latin="Calibri", east_asia="Microsoft YaHei", size=13, color=BLUE)
    h2.font.bold = True
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.line_spacing = 1.0
    h2.paragraph_format.keep_with_next = True

    h3 = doc.styles["Heading 3"]
    set_style_font(h3, latin="Calibri", east_asia="Microsoft YaHei", size=12, color=DARK_BLUE)
    h3.font.bold = True
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.line_spacing = 1.0
    h3.paragraph_format.keep_with_next = True

    caption = doc.styles["Caption"]
    set_style_font(caption, latin="Calibri", east_asia="SimSun", size=9.5, color=BLACK)
    caption.font.italic = False
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.line_spacing = 1.1
    caption.paragraph_format.keep_together = True

    if "Equation" not in [s.name for s in doc.styles]:
        equation = doc.styles.add_style("Equation", 1)
    else:
        equation = doc.styles["Equation"]
    set_style_font(equation, latin="Cambria Math", east_asia="Cambria Math", size=10.5)
    equation.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    equation.paragraph_format.space_before = Pt(6)
    equation.paragraph_format.space_after = Pt(8)
    equation.paragraph_format.keep_together = True

    if "Reference" not in [s.name for s in doc.styles]:
        reference = doc.styles.add_style("Reference", 1)
    else:
        reference = doc.styles["Reference"]
    set_style_font(reference, latin="Calibri", east_asia="SimSun", size=9.5)
    reference.paragraph_format.left_indent = Inches(0.25)
    reference.paragraph_format.first_line_indent = Inches(-0.25)
    reference.paragraph_format.space_before = Pt(0)
    reference.paragraph_format.space_after = Pt(4)
    reference.paragraph_format.line_spacing = 1.15

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hp.paragraph_format.space_after = Pt(0)
    hr = hp.add_run("专注音乐的多尺度有向拓扑指纹 · 验证研究")
    set_run_font(hr, size=8.5, color=MUTED)

    first_header = section.first_page_header
    first_header.paragraphs[0].text = ""
    footer = section.footer
    add_page_field(footer.paragraphs[0])
    section.first_page_footer.paragraphs[0].text = ""

    doc.core_properties.title = "专注音乐的多尺度有向拓扑指纹"
    doc.core_properties.subject = "GLMY 路径同调、持久同调与相位提升环结构的验证研究"
    doc.core_properties.keywords = (
        "专注音乐; GLMY 路径同调; 持久同调; 有向图; 相位提升; 拓扑数据分析"
    )
    doc.core_properties.author = ""


def add_cover(doc: Document, lines: list[str]) -> None:
    title_text = lines[0].removeprefix("# ").strip()
    subtitle_text = lines[2].removeprefix("## ").strip()
    metadata = [line.strip() for line in lines[4:] if line.strip()]

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(92)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    kr = kicker.add_run("研究论文  |  拓扑音乐分析")
    set_run_font(kr, east_asia="Microsoft YaHei", size=10.5, color=RGBColor(0x7A, 0x5A, 0x00), bold=True)

    p = doc.add_paragraph(style="Title")
    p.add_run(title_text)
    p = doc.add_paragraph(style="Subtitle")
    p.add_run(subtitle_text)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(58)

    for index, item in enumerate(metadata):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(8 if index < len(metadata) - 1 else 0)
        run = p.add_run(item)
        set_run_font(run, east_asia="Microsoft YaHei", size=11.5, color=BLACK)

    doc.add_page_break()


def set_cell_margins(cell, *, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), TABLE_BORDER)


def apply_table_geometry(table, widths_dxa: list[int]) -> None:
    if sum(widths_dxa) != CONTENT_DXA:
        raise ValueError(f"table widths must sum to {CONTENT_DXA}: {widths_dxa}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        prevent_row_split(row)
        for cell, width in zip(row.cells, widths_dxa, strict=True):
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def table_widths(headers: list[str]) -> list[int]:
    count = len(headers)
    if count == 7:
        return [820, 900, 1040, 1030, 1060, 900, 3610]
    if count == 3:
        return [1550, 5000, 2810]
    if count == 8:
        return [700, 2250, 850, 850, 720, 960, 960, 2070]
    base = CONTENT_DXA // count
    widths = [base] * count
    widths[-1] += CONTENT_DXA - sum(widths)
    return widths


def add_table(doc: Document, rows: list[list[str]]) -> None:
    headers = rows[0]
    data = rows[1:]
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_borders(table)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for j, text in enumerate(headers):
        cell = hdr.cells[j]
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), TABLE_FILL)
        cell._tc.get_or_add_tcPr().append(shading)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.1
        r = p.add_run(text)
        set_run_font(r, size=8.6 if len(headers) >= 7 else 9.2, bold=True)

    for values in data:
        row = table.add_row()
        for j, text in enumerate(values):
            p = row.cells[j].paragraphs[0]
            numeric = bool(re.fullmatch(r"[0-9.+−×^()–/ ]+", text))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if numeric or j == 0 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            r = p.add_run(text)
            set_run_font(r, size=8.4 if len(headers) >= 7 else 9.0)
    apply_table_geometry(table, table_widths(headers))
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)


def add_inline_runs(paragraph, text: str, *, size: float | None = None) -> None:
    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=size, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, latin="Consolas", east_asia="Microsoft YaHei", size=(size or 10.5) - 0.5)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=size)


def add_equation(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Equation")
    text = text.replace("i_p", "iₚ")
    pattern = re.compile(r"(_\([^)]+\)|_[A-Za-z0-9]+|\^-?[0-9]+)")
    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            run = p.add_run(text[cursor : match.start()])
            set_run_font(run, latin="Cambria Math", east_asia="Cambria Math", size=10.5)
        token = match.group(0)
        run = p.add_run(token[2:-1] if token.startswith("_(") else token[1:])
        set_run_font(run, latin="Cambria Math", east_asia="Cambria Math", size=9.5)
        if token.startswith("^"):
            run.font.superscript = True
        else:
            run.font.subscript = True
        cursor = match.end()
    if cursor < len(text):
        run = p.add_run(text[cursor:])
        set_run_font(run, latin="Cambria Math", east_asia="Cambria Math", size=10.5)


def add_figure(doc: Document, caption: str, raw_path: str) -> None:
    image_path = (SOURCE.parent / raw_path).resolve()
    if not image_path.is_file():
        raise FileNotFoundError(image_path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    width = 6.35
    if "topology_pca" in image_path.name:
        width = 5.65
    elif "classification" in image_path.name:
        width = 5.85
    run = p.add_run()
    inline = run.add_picture(str(image_path), width=Inches(width))
    inline._inline.docPr.set("descr", caption)
    cp = doc.add_paragraph(style="Caption")
    add_inline_runs(cp, caption, size=9.5)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    raw: list[str] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        raw.append(lines[i].strip())
        i += 1
    rows = [[cell.strip() for cell in row.strip("|").split("|")] for row in raw]
    if len(rows) < 2:
        raise ValueError("malformed markdown table")
    separator = rows[1]
    if not all(re.fullmatch(r":?-{3,}:?", cell) for cell in separator):
        raise ValueError("markdown table separator missing")
    return [rows[0], *rows[2:]], i


def add_body(doc: Document, lines: list[str]) -> None:
    i = 0
    in_references = False
    while i < len(lines):
        raw = lines[i]
        text = raw.strip()
        if not text:
            i += 1
            continue
        if text == "<!-- PAGEBREAK -->":
            doc.add_page_break()
            i += 1
            continue
        if text.startswith("# "):
            heading = text[2:].strip()
            p = doc.add_paragraph(heading, style="Heading 1")
            in_references = heading == "参考文献"
            i += 1
            continue
        if text.startswith("## "):
            doc.add_paragraph(text[3:].strip(), style="Heading 2")
            i += 1
            continue
        if text.startswith("### "):
            doc.add_paragraph(text[4:].strip(), style="Heading 3")
            i += 1
            continue
        if text == "$$":
            formula: list[str] = []
            i += 1
            while i < len(lines) and lines[i].strip() != "$$":
                formula.append(lines[i].strip())
                i += 1
            add_equation(doc, " ".join(formula))
            i += 1
            continue
        image_match = re.fullmatch(r"!\[(.+)]\((.+)\)", text)
        if image_match:
            add_figure(doc, image_match.group(1), image_match.group(2))
            i += 1
            continue
        if text.startswith("|"):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue
        if re.match(r"^表\s*\d+", text):
            p = doc.add_paragraph(style="Caption")
            p.paragraph_format.keep_with_next = True
            p.paragraph_format.space_after = Pt(4)
            add_inline_runs(p, text, size=9.5)
            i += 1
            continue
        if text.startswith("注："):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.line_spacing = 1.15
            add_inline_runs(p, text, size=9.2)
            for run in p.runs:
                run.font.color.rgb = MUTED
            i += 1
            continue
        style = "Reference" if in_references and not text.startswith("附录") else "Normal"
        p = doc.add_paragraph(style=style)
        add_inline_runs(p, text)
        i += 1


def preset_audit(doc: Document) -> None:
    section = doc.sections[0]
    assert section.page_width == Inches(8.5)
    assert section.page_height == Inches(11)
    assert section.left_margin == Inches(1)
    assert section.right_margin == Inches(1)
    assert doc.styles["Normal"].font.size == Pt(11)
    for table in doc.tables:
        tbl_w = table._tbl.tblPr.first_child_found_in("w:tblW")
        tbl_ind = table._tbl.tblPr.first_child_found_in("w:tblInd")
        assert tbl_w is not None and tbl_w.get(qn("w:w")) == str(CONTENT_DXA)
        assert tbl_ind is not None and tbl_ind.get(qn("w:w")) == str(TABLE_INDENT_DXA)
        grid_sum = sum(int(col.get(qn("w:w"))) for col in table._tbl.tblGrid)
        assert grid_sum == CONTENT_DXA


def main() -> int:
    parser = argparse.ArgumentParser(description="Build the focus-music topology paper DOCX.")
    parser.add_argument(
        "--output",
        type=Path,
        default=OUTPUT,
        help="Output DOCX path (defaults to the canonical file under output/).",
    )
    args = parser.parse_args()

    text = SOURCE.read_text(encoding="utf-8").replace("\u2011", "-")
    lines = text.splitlines()
    marker = lines.index("<!-- PAGEBREAK -->")
    cover_lines = lines[:marker]
    body_lines = lines[marker + 1 :]

    doc = Document()
    configure_document(doc)
    add_cover(doc, cover_lines)
    add_body(doc, body_lines)
    preset_audit(doc)

    output = args.output.resolve()
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    print(output)
    print(f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} inline_shapes={len(doc.inline_shapes)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
