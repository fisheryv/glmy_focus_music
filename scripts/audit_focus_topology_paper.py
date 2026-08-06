from __future__ import annotations

import argparse
import json
import re
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "output" / "专注音乐多尺度有向拓扑指纹_完整论文.docx"


def main() -> int:
    parser = argparse.ArgumentParser(description="Audit the generated focus-music topology paper DOCX.")
    parser.add_argument(
        "docx",
        nargs="?",
        type=Path,
        default=DOCX,
        help="DOCX to audit (defaults to the canonical file under output/).",
    )
    args = parser.parse_args()
    docx_path = args.docx.resolve()
    doc = Document(docx_path)
    checks: dict[str, object] = {}
    section = doc.sections[0]
    checks["page_size_twips"] = [section.page_width.twips, section.page_height.twips]
    checks["margins_twips"] = [
        section.top_margin.twips,
        section.right_margin.twips,
        section.bottom_margin.twips,
        section.left_margin.twips,
    ]
    checks["paragraphs"] = len(doc.paragraphs)
    checks["tables"] = len(doc.tables)
    checks["inline_shapes"] = len(doc.inline_shapes)

    table_geometry = []
    for table in doc.tables:
        tbl_pr = table._tbl.tblPr
        tbl_w = tbl_pr.first_child_found_in("w:tblW")
        tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
        grid = [int(col.get(qn("w:w"))) for col in table._tbl.tblGrid]
        header = table.rows[0]._tr.get_or_add_trPr().find(qn("w:tblHeader")) is not None
        table_geometry.append(
            {
                "width": int(tbl_w.get(qn("w:w"))),
                "indent": int(tbl_ind.get(qn("w:w"))),
                "grid_sum": sum(grid),
                "columns": len(grid),
                "header_repeat": header,
                "empty_cells": sum(not cell.text.strip() for row in table.rows for cell in row.cells),
            }
        )
    checks["table_geometry"] = table_geometry

    with zipfile.ZipFile(docx_path) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
        styles_xml = archive.read("word/styles.xml").decode("utf-8")
        rels_xml = archive.read("word/_rels/document.xml.rels").decode("utf-8")
        media = [name for name in archive.namelist() if name.startswith("word/media/")]
    forbidden = ["待填", "预期结果", "tool citation", "turn0", "Spotify Pop"]
    forbidden.extend(
        [
            "1,598",
            "4,794",
            "pseudo-F=2.8848",
            "Macro-F1=0.754",
            "ace_rerank_180s_v1",
        ]
    )
    checks["forbidden_hits"] = [token for token in forbidden if token in document_xml]
    required = [
        "800 首独立器乐音频",
        "pseudo-F=3.1425",
        "q=0.777",
        "不支持",
        "1.82×10^-6",
        "0.00194",
        "Macro-F1=0.776",
    ]
    checks["missing_required_text"] = [token for token in required if token not in document_xml]
    wp_doc_pr = re.findall(r"<wp:docPr[^>]+>", document_xml)
    checks["image_alt_count"] = sum(
        'descr="' in node and 'descr=""' not in node for node in wp_doc_pr
    )
    checks["media_files"] = len(media)
    checks["image_relationships"] = rels_xml.count("relationships/image")
    checks["styles_have_heading_levels"] = all(
        token in styles_xml for token in ("Heading1", "Heading2", "Heading3")
    )

    failures = []
    if checks["page_size_twips"] != [12240, 15840]:
        failures.append("page size")
    if checks["margins_twips"] != [1440, 1440, 1440, 1440]:
        failures.append("margins")
    if checks["tables"] != 3 or checks["inline_shapes"] != 4:
        failures.append("artifact count")
    if any(
        item["width"] != 9360
        or item["indent"] != 120
        or item["grid_sum"] != 9360
        or not item["header_repeat"]
        or item["empty_cells"]
        for item in table_geometry
    ):
        failures.append("table geometry")
    if checks["forbidden_hits"] or checks["missing_required_text"]:
        failures.append("content")
    if checks["image_alt_count"] != 4 or checks["media_files"] != 4 or checks["image_relationships"] != 4:
        failures.append("images")
    if not checks["styles_have_heading_levels"]:
        failures.append("styles")
    checks["failures"] = failures
    print(json.dumps(checks, ensure_ascii=False, indent=2))
    return 1 if failures else 0


if __name__ == "__main__":
    raise SystemExit(main())
